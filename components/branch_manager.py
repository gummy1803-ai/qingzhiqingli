"""分支管理系统: 实现文件系统的分支隔离、版本控制和协作功能。

核心功能:
1. 分支管理: 创建/切换/删除/重命名分支
2. 文件快照: 记录每个分支的文件状态
3. 重复性检测: 基于 SHA256 哈希检测重复文件
4. 完整性校验: 检测损坏或无法读取的文件
5. 合并与冲突解决: 支持分支合并和冲突处理
"""

import hashlib
import logging
import os
import shutil
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

logger = logging.getLogger(__name__)


class BranchManager:
    """分支管理器核心类。"""
    
    def __init__(self, base_path: str = None, branches_dir: str = None):
        """初始化分支管理器。

        支持三种方式指定根目录（优先级从高到低）:
          1. 构造函数显式参数 (base_path 或兼容别名 branches_dir)
          2. 环境变量 BRANCH_ROOT_DIR (用于 CI / 冒烟测试隔离)
          3. 默认: 项目根目录下的 .branches

        Args:
            base_path: 分支存储的基础路径
            branches_dir: 兼容别名, 与 base_path 等价 (冒烟脚本使用)
        """
        explicit = base_path or branches_dir
        if explicit:
            base_path = explicit
        elif os.environ.get("BRANCH_ROOT_DIR", "").strip():
            base_path = os.environ["BRANCH_ROOT_DIR"].strip()
        else:
            base_path = Path(__file__).parent.parent / ".branches"
        self.base_path = Path(base_path)
        self.base_path.mkdir(parents=True, exist_ok=True)
        logger.info("[分支管理] 初始化: 基础路径=%s", self.base_path)
        self._initialize_default_branch()
    
    def _initialize_default_branch(self):
        """初始化默认主分支。"""
        main_branch = self.base_path / "main"
        if not main_branch.exists():
            main_branch.mkdir(parents=True, exist_ok=True)
            self._create_branch_metadata("main", "主分支 (默认)", is_active=True)
    
    def _get_branch_path(self, branch_name: str) -> Path:
        """获取分支的文件存储路径。"""
        return self.base_path / branch_name
    
    def _get_metadata_path(self, branch_name: str) -> Path:
        """获取分支元数据文件路径。"""
        return self.base_path / f"{branch_name}.meta.json"
    
    def _create_branch_metadata(
        self,
        branch_name: str,
        description: str = "",
        is_active: bool = False,
        parent_branch: str = None
    ):
        """创建分支元数据。"""
        import json
        metadata = {
            "name": branch_name,
            "description": description,
            "is_active": is_active,
            "parent_branch": parent_branch,
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat(),
            "versions": [],
            "conflicts": []
        }
        meta_path = self._get_metadata_path(branch_name)
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(metadata, f, ensure_ascii=False, indent=2)
    
    def _read_branch_metadata(self, branch_name: str) -> Dict:
        """读取分支元数据。"""
        import json
        meta_path = self._get_metadata_path(branch_name)
        if not meta_path.exists():
            return {}
        with open(meta_path, "r", encoding="utf-8") as f:
            return json.load(f)
    
    def _update_branch_metadata(self, branch_name: str, updates: Dict):
        """更新分支元数据。"""
        import json
        meta = self._read_branch_metadata(branch_name)
        meta.update(updates)
        meta["updated_at"] = datetime.now().isoformat()
        meta_path = self._get_metadata_path(branch_name)
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
    
    def list_branches(self) -> List[Dict]:
        """列出所有分支。
        
        注意: Path.stem 仅剥去最后一个扩展名, 对 'xxx.meta.json' 会得到 'xxx.meta'。
        因此需用 .stem.stem 或 remove_suffix('.meta') 获取真正的分支名。
        """
        branches = []
        for meta_file in self.base_path.glob("*.meta.json"):
            # 移除双重扩展名: main.meta.json → main
            branch_name = meta_file.name[:-len(".meta.json")] if meta_file.name.endswith(".meta.json") else meta_file.stem
            meta = self._read_branch_metadata(branch_name)
            if meta:
                branch_path = self._get_branch_path(branch_name)
                file_count = sum(1 for f in branch_path.rglob("*") if f.is_file())
                total_size = sum(f.stat().st_size for f in branch_path.rglob("*") if f.is_file())
                meta["file_count"] = file_count
                meta["total_size"] = total_size
                branches.append(meta)
        return sorted(branches, key=lambda x: x.get("name", ""))
    
    def create_branch(
        self,
        branch_name: str,
        description: str = "",
        source_branch: str = None
    ) -> Tuple[bool, str]:
        """创建新分支。
        
        Args:
            branch_name: 新分支名称
            description: 分支描述
            source_branch: 来源分支（fork时使用）
            
        Returns:
            (成功标志, 消息)
        """
        t0 = time.perf_counter()
        logger.info("[分支创建] 请求: name=%s, source=%s, desc=%s", 
                    branch_name, source_branch, description[:50])
        
        if not branch_name or not branch_name.replace("-", "").replace("_", "").isalnum():
            logger.warning("[分支创建] 失败: 分支名无效 '%s'", branch_name)
            return False, "分支名只能包含字母、数字、-、_ 字符"
        
        branch_path = self._get_branch_path(branch_name)
        if branch_path.exists():
            logger.warning("[分支创建] 失败: 分支 '%s' 已存在", branch_name)
            return False, f"分支 '{branch_name}' 已存在"
        
        try:
            if source_branch:
                # 从源分支复制文件
                source_path = self._get_branch_path(source_branch)
                if not source_path.exists():
                    logger.error("[分支创建] 失败: 源分支 '%s' 不存在", source_branch)
                    return False, f"源分支 '{source_branch}' 不存在"
                
                file_count = sum(1 for f in source_path.rglob("*") if f.is_file())
                logger.info("[分支创建] Fork: 从 '%s' 复制 %d 个文件到 '%s'", 
                           source_branch, file_count, branch_name)
                shutil.copytree(source_path, branch_path)
                self._create_branch_metadata(
                    branch_name,
                    description,
                    is_active=False,
                    parent_branch=source_branch
                )
                self._add_version_record(branch_name, "create", 
                    f"从分支 '{source_branch}' 创建", file_count)
                elapsed = (time.perf_counter() - t0) * 1000
                logger.info("[分支创建] 成功: '%s' (fork自'%s', %d文件, %.1fms)", 
                           branch_name, source_branch, file_count, elapsed)
                return True, f"成功创建分支 '{branch_name}' (来源于 '{source_branch}')"
            else:
                # 创建空分支
                branch_path.mkdir(parents=True, exist_ok=True)
                self._create_branch_metadata(
                    branch_name, description, is_active=False
                )
                self._add_version_record(branch_name, "create", "创建空分支", 0)
                elapsed = (time.perf_counter() - t0) * 1000
                logger.info("[分支创建] 成功: 空分支 '%s' (%.1fms)", branch_name, elapsed)
                return True, f"成功创建空分支 '{branch_name}'"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[分支创建] 异常: '%s' -> %s (%.1fms)", branch_name, str(e), elapsed)
            return False, f"创建分支失败: {str(e)}"
    
    def switch_branch(self, branch_name: str) -> Tuple[bool, str]:
        """切换到指定分支。"""
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            return False, f"分支 '{branch_name}' 不存在"
        
        # 将所有分支设为非活跃
        for b in self.list_branches():
            self._update_branch_metadata(b["name"], {"is_active": False})
        
        # 激活目标分支
        self._update_branch_metadata(branch_name, {"is_active": True})
        return True, f"已切换到分支 '{branch_name}'"
    
    def get_active_branch(self) -> Optional[str]:
        """获取当前活跃分支。"""
        for b in self.list_branches():
            if b.get("is_active"):
                return b["name"]
        return "main"
    
    def rename_branch(self, old_name: str, new_name: str) -> Tuple[bool, str]:
        """重命名分支。"""
        old_path = self._get_branch_path(old_name)
        new_path = self._get_branch_path(new_name)
        
        if not old_path.exists():
            return False, f"分支 '{old_name}' 不存在"
        if new_path.exists():
            return False, f"分支 '{new_name}' 已存在"
        
        try:
            # 移动文件目录
            old_path.rename(new_path)
            
            # 移动元数据文件
            old_meta = self._get_metadata_path(old_name)
            new_meta = self._get_metadata_path(new_name)
            if old_meta.exists():
                meta = self._read_branch_metadata(old_name)
                meta["name"] = new_name
                with open(new_meta, "w", encoding="utf-8") as f:
                    import json
                    json.dump(meta, f, ensure_ascii=False, indent=2)
                old_meta.unlink()
            
            self._add_version_record(new_name, "rename", 
                f"从 '{old_name}' 重命名", 0)
            return True, f"分支已重命名为 '{new_name}'"
        except Exception as e:
            return False, f"重命名失败: {str(e)}"
    
    def delete_branch(self, branch_name: str, force: bool = False) -> Tuple[bool, str]:
        """删除分支。
        
        Args:
            branch_name: 要删除的分支名
            force: 强制删除（即使有文件）
        """
        if branch_name == "main":
            return False, "不能删除主分支 'main'"
        
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            return False, f"分支 '{branch_name}' 不存在"
        
        file_count = sum(1 for f in branch_path.rglob("*") if f.is_file())
        if file_count > 0 and not force:
            return False, f"分支 '{branch_name}' 包含 {file_count} 个文件, 使用 force=True 强制删除"
        
        try:
            shutil.rmtree(branch_path)
            meta_path = self._get_metadata_path(branch_name)
            if meta_path.exists():
                meta_path.unlink()
            return True, f"分支 '{branch_name}' 已删除"
        except Exception as e:
            return False, f"删除失败: {str(e)}"

    # ====================================================================
    # 文件夹管理 (Folder CRUD)
    # ====================================================================

    def list_folders(self, branch_name: str, relative_to: str = "") -> List[Dict]:
        """列出分支中的所有文件夹（默认平铺；指定relative_to可只看某子目录下一层）。

        返回: [{"path": "raw/车212", "name": "车212", "parent": "raw", "file_count": 3, "size": 10240}, ...]
        """
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            logger.warning("[文件夹列表] 分支不存在: %s", branch_name)
            return []

        root = branch_path / relative_to if relative_to else branch_path
        if not root.exists():
            logger.warning("[文件夹列表] 相对路径不存在: branch=%s, rel=%s", branch_name, relative_to)
            return []

        results: List[Dict] = []
        # 收集所有目录
        for dir_path in sorted(root.rglob("*")):
            if not dir_path.is_dir():
                continue
            rel = str(dir_path.relative_to(branch_path)).replace("\\", "/")
            if rel == ".":
                continue
            # 文件数 = 该目录下（含子目录）所有文件数
            files_under = [p for p in dir_path.rglob("*") if p.is_file()]
            size_under = sum(p.stat().st_size for p in files_under)
            parent = str(Path(rel).parent).replace("\\", "/")
            if parent == ".":
                parent = ""
            results.append({
                "path": rel,
                "name": dir_path.name,
                "parent": parent,
                "file_count": len(files_under),
                "size": size_under,
            })
        elapsed = (time.perf_counter() - t0) * 1000
        logger.info("[文件夹列表] branch=%s, scope=%s, 共 %d 个目录 (%.1fms)",
                    branch_name, relative_to or "(全分支)", len(results), elapsed)
        return results

    def create_folder(self, branch_name: str, folder_path: str) -> Tuple[bool, str]:
        """在分支中创建文件夹（支持多级，如 'raw/车212/第一批次'）。"""
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            logger.error("[文件夹创建] 分支不存在: %s", branch_name)
            return False, f"分支 '{branch_name}' 不存在"

        if not folder_path or not folder_path.strip():
            return False, "文件夹路径不能为空"

        # 规范化 + 防穿越
        fp_norm = folder_path.strip().replace("\\", "/").strip("/")
        target = (branch_path / fp_norm).resolve()
        try:
            target.relative_to(branch_path.resolve())
        except ValueError:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹创建] 路径越界: branch=%s, path=%s (%.1fms)",
                           branch_name, folder_path, elapsed)
            return False, "文件夹路径不能跳出分支目录"

        if target.exists() and target.is_dir():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件夹创建] 目录已存在，跳过: branch=%s, path=%s (%.1fms)",
                        branch_name, fp_norm, elapsed)
            return True, f"文件夹 '{fp_norm}' 已存在"
        if target.exists() and target.is_file():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹创建] 同名文件已存在: branch=%s, path=%s (%.1fms)",
                           branch_name, fp_norm, elapsed)
            return False, f"同名文件 '{fp_norm}' 已存在，请先删除"

        try:
            target.mkdir(parents=True, exist_ok=True)
            self._add_version_record(branch_name, "commit", f"新建文件夹: {fp_norm}", -1)
            self._sync_snapshot_to_db(branch_name, change_note=f"mkdir:{fp_norm}")
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件夹创建] 成功: branch=%s, path=%s (%.1fms)",
                        branch_name, fp_norm, elapsed)
            return True, f"已创建文件夹 '{fp_norm}'"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件夹创建] 异常: branch=%s, path=%s, err=%s (%.1fms)",
                        branch_name, folder_path, str(e), elapsed, exc_info=True)
            return False, f"创建失败: {str(e)}"

    def rename_folder(self, branch_name: str, old_path: str, new_path: str) -> Tuple[bool, str]:
        """重命名/移动文件夹（跨目录移动也支持）。"""
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            return False, f"分支 '{branch_name}' 不存在"

        old_norm = old_path.strip().replace("\\", "/").strip("/")
        new_norm = new_path.strip().replace("\\", "/").strip("/")
        if not old_norm or not new_norm:
            return False, "旧路径和新路径都不能为空"

        old_target = (branch_path / old_norm).resolve()
        new_target = (branch_path / new_norm).resolve()
        try:
            old_target.relative_to(branch_path.resolve())
            new_target.relative_to(branch_path.resolve())
        except ValueError:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹重命名] 路径越界: branch=%s, old=%s, new=%s (%.1fms)",
                           branch_name, old_path, new_path, elapsed)
            return False, "路径不能跳出分支目录"

        if not old_target.exists() or not old_target.is_dir():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹重命名] 旧目录不存在: branch=%s, old=%s (%.1fms)",
                           branch_name, old_norm, elapsed)
            return False, f"旧文件夹 '{old_norm}' 不存在"

        if new_target.exists():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹重命名] 目标已存在: branch=%s, new=%s (%.1fms)",
                           branch_name, new_norm, elapsed)
            return False, f"目标路径 '{new_norm}' 已存在"

        # 防嵌套：new_path 不能是 old_path 的子路径（否则会自己套自己形成循环）
        try:
            Path(new_target).resolve().relative_to(old_target.resolve())
            return False, f"新路径 '{new_norm}' 是旧路径 '{old_norm}' 的子目录，禁止操作"
        except ValueError:
            pass

        try:
            # 确保父目录存在
            new_target.parent.mkdir(parents=True, exist_ok=True)
            old_target.rename(new_target)
            self._add_version_record(branch_name, "commit",
                f"重命名文件夹: {old_norm} -> {new_norm}", -1)
            self._sync_snapshot_to_db(branch_name, change_note=f"mvdir:{old_norm}->{new_norm}")
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件夹重命名] 成功: branch=%s, %s -> %s (%.1fms)",
                        branch_name, old_norm, new_norm, elapsed)
            return True, f"已重命名文件夹: '{old_norm}' → '{new_norm}'"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件夹重命名] 异常: branch=%s, old=%s, new=%s, err=%s (%.1fms)",
                        branch_name, old_norm, new_norm, str(e), elapsed, exc_info=True)
            return False, f"重命名失败: {str(e)}"

    def delete_folder(self, branch_name: str, folder_path: str, force: bool = False) -> Tuple[bool, str]:
        """删除文件夹（默认非空需要 force=True，防误删）。"""
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            return False, f"分支 '{branch_name}' 不存在"

        norm = folder_path.strip().replace("\\", "/").strip("/")
        if not norm:
            return False, "文件夹路径不能为空"
        target = (branch_path / norm).resolve()
        try:
            target.relative_to(branch_path.resolve())
        except ValueError:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹删除] 路径越界: branch=%s, path=%s (%.1fms)",
                           branch_name, folder_path, elapsed)
            return False, "路径不能跳出分支目录"

        if not target.exists() or not target.is_dir():
            return False, f"文件夹 '{norm}' 不存在"

        files_under = [p for p in target.rglob("*") if p.is_file()]
        if files_under and not force:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件夹删除] 目录非空且未force: branch=%s, path=%s, files=%d (%.1fms)",
                           branch_name, norm, len(files_under), elapsed)
            return False, f"文件夹 '{norm}' 内有 {len(files_under)} 个文件，请勾选 force=True 再删除"

        try:
            shutil.rmtree(target)
            self._add_version_record(branch_name, "commit",
                f"删除文件夹: {norm} (files={len(files_under)})", -1)
            self._sync_snapshot_to_db(branch_name, change_note=f"rmdir:{norm}")
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件夹删除] 成功: branch=%s, path=%s, removed_files=%d (%.1fms)",
                        branch_name, norm, len(files_under), elapsed)
            return True, f"已删除文件夹 '{norm}'（含 {len(files_under)} 个文件）"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件夹删除] 异常: branch=%s, path=%s, err=%s (%.1fms)",
                        branch_name, norm, str(e), elapsed, exc_info=True)
            return False, f"删除失败: {str(e)}"

    # ====================================================================
    # 文件管理增强：指定目录导入、移动、预览打开
    # ====================================================================

    def add_file_to_folder(
        self,
        branch_name: str,
        file_path: str,
        file_content: Optional[bytes] = None,
        target_folder: str = "",
        target_name: Optional[str] = None,
        upload_user: str = "ui",
    ) -> Tuple[bool, str, str]:
        """导入文件并存入指定子文件夹（核心入口，UI应使用此函数而非add_file）。

        此函数在文件系统写入成功后会触发：
        1) branch_file_snapshots 表同步（一致性）
        2) 如果是 CSV / DOCX / XLSX 等业务数据 → 写入 vehicle_data_files 业务表
        """
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            logger.error("[文件导入] 分支不存在: branch=%s", branch_name)
            return False, f"分支 '{branch_name}' 不存在", ""

        folder_norm = (target_folder or "").strip().replace("\\", "/").strip("/")
        if target_name is None:
            target_name = Path(file_path).name
        rel_combined = (Path(folder_norm) / target_name).as_posix() if folder_norm else target_name

        logger.info("[文件导入] 开始: branch=%s, src=%s, folder=%s, name=%s",
                    branch_name, Path(file_path).name if file_path else "(bytes)",
                    folder_norm or "/(根)", target_name)

        # 直接复用 add_file 的 target_name 能力（现在 target_name 可以带相对路径）
        success, msg, file_hash = self.add_file(branch_name, file_path, file_content,
                                                target_name=rel_combined)
        if not success:
            return False, msg, ""

        # ---- 写入一致性/业务库（钩子） ----
        stored_path = branch_path / rel_combined
        size_kb = stored_path.stat().st_size // 1024 if stored_path.exists() else 0
        try:
            # 1) DB 快照一致性同步
            self._sync_snapshot_to_db(branch_name, change_note=f"import:{rel_combined}")
            # 2) 业务库落库（CSV/docx/xlsx → vehicle_data_files）
            data_kind, vehicle_id = self._persist_business_file(
                branch_name, rel_combined, stored_path, upload_user=upload_user)
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info(
                "[文件导入] ✅ 完成: branch=%s, file=%s, size=%dKB, hash=%s…, "
                "data_kind=%s, vehicle=%s (%.1fms)",
                branch_name, rel_combined, size_kb, (file_hash or "")[:16],
                data_kind or "(未识别)", vehicle_id or "(空)", elapsed,
            )
            return True, msg + f"（类型:{data_kind or '未识别'}, 车号:{vehicle_id or '-'})", file_hash
        except Exception as sync_e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件导入] 文件系统已写入，但DB同步失败: branch=%s, file=%s, err=%s (%.1fms)",
                        branch_name, rel_combined, str(sync_e), elapsed, exc_info=True)
            # 文件系统已经成功，仍向调用方返回成功但附带提示
            return True, msg + f"（⚠️ DB同步失败：{sync_e}）", file_hash

    def move_file(self, branch_name: str, src_path: str, dst_path: str) -> Tuple[bool, str]:
        """移动文件（跨目录），等价于重命名文件，语义独立出来便于UI理解。"""
        import time
        t0 = time.perf_counter()
        logger.info("[文件移动] 开始: branch=%s, %s -> %s", branch_name, src_path, dst_path)
        ok, msg = self.rename_file(branch_name, src_path, dst_path)
        elapsed = (time.perf_counter() - t0) * 1000
        if ok:
            logger.info("[文件移动] 成功 (%.1fms)", elapsed)
        else:
            logger.warning("[文件移动] 失败: %s (%.1fms)", msg, elapsed)
        return ok, msg

    def open_file_preview(
        self,
        branch_name: str,
        file_path: str,
        max_rows: int = 200,
    ) -> Dict:
        """打开并预览分支内的文件（返回统一结构，便于Streamlit渲染）。

        返回: {"ok": bool, "error": str, "kind": "csv|excel|docx|text|binary|image",
               "rows": int, "cols": int, "columns": [...],
               "dataframe": pd.DataFrame | None, "text": str,
               "raw_path": str, "size": int}
        """
        import time, io
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        target = branch_path / file_path
        result: Dict = {"ok": False, "error": "", "kind": "binary", "rows": 0, "cols": 0,
                        "columns": [], "dataframe": None, "text": "",
                        "raw_path": str(target), "size": 0}

        if not target.exists():
            result["error"] = f"文件 '{file_path}' 不存在"
            logger.warning("[文件预览] 文件不存在: branch=%s, path=%s", branch_name, file_path)
            return result

        result["size"] = target.stat().st_size
        suffix = target.suffix.lower()
        try:
            # --- CSV / TSV ---
            if suffix in (".csv", ".txt"):
                import pandas as _pd
                try:
                    df = _pd.read_csv(target, nrows=max_rows)
                except UnicodeDecodeError:
                    df = _pd.read_csv(target, nrows=max_rows, encoding="gbk")
                result.update({
                    "ok": True, "kind": "csv",
                    "rows": len(df), "cols": len(df.columns),
                    "columns": list(df.columns), "dataframe": df,
                })
            # --- Excel ---
            elif suffix in (".xlsx", ".xls"):
                import pandas as _pd
                df = _pd.read_excel(target, nrows=max_rows)
                result.update({
                    "ok": True, "kind": "excel",
                    "rows": len(df), "cols": len(df.columns),
                    "columns": list(df.columns), "dataframe": df,
                })
            # --- DOCX ---
            elif suffix == ".docx":
                try:
                    from docx import Document  # type: ignore
                    doc = Document(str(target))
                    paras = [p.text for p in doc.paragraphs]
                    # 还尝试取第一个表格（耐久工步docx常含表格）
                    tables_text = []
                    for ti, tbl in enumerate(doc.tables[:3]):
                        head = [c.text.strip() for c in tbl.rows[0].cells] if tbl.rows else []
                        rows2d = [[c.text.strip() for c in r.cells] for r in tbl.rows[1:101]]
                        import pandas as _pd
                        df_tbl = _pd.DataFrame(rows2d, columns=head[:len(rows2d[0])] if rows2d else head)
                        tables_text.append(df_tbl)
                    result.update({
                        "ok": True, "kind": "docx",
                        "text": "\n".join(paras[:200]),
                        "dataframe": tables_text[0] if tables_text else None,
                        "rows": sum(len(t) for t in tables_text) if tables_text else len(paras),
                        "columns": list(tables_text[0].columns) if tables_text else [],
                        "tables": tables_text,
                    })
                except Exception as de:
                    result["error"] = f"DOCX解析失败: {de}"
            # --- 图片 ---
            elif suffix in (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"):
                result.update({"ok": True, "kind": "image"})
            # --- 默认按文本/二进制兜底 ---
            else:
                raw = target.read_bytes()
                try:
                    result["text"] = raw.decode("utf-8", errors="replace")[:100_000]
                    result["ok"] = True
                    result["kind"] = "text"
                except Exception:
                    result["ok"] = True
                    result["kind"] = "binary"

            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件预览] %s branch=%s, path=%s, size=%dB, kind=%s (%.1fms)",
                        "✅" if result["ok"] else "⚠️", branch_name, file_path,
                        result["size"], result["kind"], elapsed)
            return result
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            result["error"] = f"预览异常: {str(e)}"
            logger.error("[文件预览] 异常: branch=%s, path=%s, err=%s (%.1fms)",
                        branch_name, file_path, str(e), elapsed, exc_info=True)
            return result

    # ====================================================================
    # 一致性同步 & 业务库落库钩子（DB 适配层：可无 DB 安全降级）
    # ====================================================================

    def _sync_snapshot_to_db(self, branch_name: str, change_note: str = "") -> None:
        """把当前分支内所有文件的快照同步写入 database.py 的 branch_file_snapshots 表。

        DB 不可用自动跳过；所有异常被吞掉只打 ERROR 日志（绝不影响文件系统主流程）。
        """
        import time
        t0 = time.perf_counter()
        try:
            from durability.database import db_sync_branch_snapshot  # 延迟import防循环
            files = self.list_branch_files(branch_name)
            ok, msg = db_sync_branch_snapshot(branch_name, files, note=change_note)
            elapsed = (time.perf_counter() - t0) * 1000
            if ok:
                logger.info("[一致性·快照] ✅ branch=%s, files=%d, note=%s (%.1fms)",
                            branch_name, len(files), change_note or "(无)", elapsed)
            else:
                logger.warning("[一致性·快照] ⚠️ branch=%s, msg=%s (%.1fms)",
                               branch_name, msg, elapsed)
        except ImportError:
            # database 模块本身缺失也不影响
            logger.debug("[一致性·快照] 跳过: durability.database 未就绪")
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[一致性·快照] ❌ branch=%s, note=%s, err=%s (%.1fms)",
                        branch_name, change_note, str(e), elapsed, exc_info=True)

    def _persist_business_file(
        self, branch_name: str, rel_path: str, stored_path: Path,
        upload_user: str = "ui",
    ) -> Tuple[Optional[str], Optional[str]]:
        """若上传的是 CSV / DOCX / XLSX 业务数据，解析基本元信息并写入 vehicle_data_files。

        返回 (data_kind, vehicle_id)，供上层日志显示。
        """
        import time, io
        t0 = time.perf_counter()
        suffix = stored_path.suffix.lower()
        data_kind: Optional[str] = None
        vehicle_id: Optional[str] = ""
        if suffix not in (".csv", ".xlsx", ".xls", ".docx"):
            return None, None
        try:
            from durability.database import (
                db_upsert_data_file, _parse_csv_filename,
            )  # 延迟import防循环
        except Exception:
            return None, None

        try:
            file_bytes = stored_path.read_bytes()
            file_hash = self._compute_file_hash(stored_path)
            row_count = 0
            col_signals: List[str] = []
            time_min: Optional[str] = None
            time_max: Optional[str] = None
            try:
                import pandas as _pd
                if suffix == ".csv":
                    df = _pd.read_csv(io.BytesIO(file_bytes), nrows=10_000)
                elif suffix in (".xlsx", ".xls"):
                    df = _pd.read_excel(io.BytesIO(file_bytes), nrows=10_000)
                else:
                    df = None
                if df is not None:
                    row_count = len(df)
                    col_signals = [str(c) for c in df.columns]
                    # 自动识别整车数据：含 Timestamp 列
                    ts_cols = [c for c in df.columns if str(c).lower() in ("timestamp", "时间", "time")]
                    if ts_cols and row_count > 0:
                        import pandas as _pd2
                        ts_series = _pd2.to_datetime(df[ts_cols[0]], errors="coerce").dropna()
                        if len(ts_series) > 0:
                            time_min = str(ts_series.min())
                            time_max = str(ts_series.max())
            except Exception as parse_e:
                logger.debug("[落库·业务] 轻量解析失败(不影响主流程): %s", parse_e)

            # --- 判断 data_kind ---
            col_set = {str(c) for c in col_signals}
            if suffix == ".docx":
                data_kind = "耐久"
            else:
                kw_cycle = any(("循环" in c or "cycle" in c.lower()) for c in col_signals)
                kw_power = any(("功率" in c or "power" in c.lower()) for c in col_signals)
                kw_stage = any(("工步" in c or "stage" in c.lower() or "avg" in c.lower()
                                or "单体电压" in c or "平均" in c) for c in col_signals)
                if kw_cycle and kw_power:
                    data_kind = "台架"
                elif "Timestamp" in col_set or "timestamp" in col_set or ts_cols:
                    data_kind = "整车"
                elif kw_stage:
                    data_kind = "耐久"
                else:
                    data_kind = "未分类"

            # --- vehicle_id: 优先用文件名解析，其次用内容猜 ---
            try:
                meta = _parse_csv_filename(stored_path.name)
            except Exception:
                meta = None
            if meta and isinstance(meta, dict) and meta.get("vehicle"):
                vehicle_id = str(meta["vehicle"])
            else:
                # 兜底：从 rel_path 里取数字段
                import re
                m = re.search(r"\d{3,}", rel_path)
                if m:
                    vehicle_id = m.group(0)

            logger.info("[落库·业务] 解析: branch=%s, file=%s, kind=%s, vehicle=%s, rows=%d, cols=%d",
                        branch_name, rel_path, data_kind, vehicle_id, row_count, len(col_signals))

            _fid, inserted, _ = db_upsert_data_file(
                data_kind=data_kind or "未分类",
                file_name=stored_path.name,
                file_bytes=file_bytes,
                file_hash=file_hash,
                vehicle_id=vehicle_id or "",
                row_count=row_count,
                time_min=time_min,
                time_max=time_max,
                col_signals=col_signals,
                upload_user=upload_user,
                status="uploaded",
                status_note=f"来源:分支[{branch_name}] 相对路径:{rel_path}",
                agg_rows=0,
                extra_meta={
                    "branch": branch_name,
                    "branch_relative_path": rel_path,
                    "imported_by": "BranchManager.add_file_to_folder",
                },
            )
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[落库·业务] %s branch=%s, file=%s, kind=%s, vehicle=%s, fid?=%s (%.1fms)",
                        "✅新写入" if inserted else "↻已存在跳过",
                        branch_name, rel_path, data_kind, vehicle_id, _fid, elapsed)
            return data_kind, vehicle_id
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[落库·业务] ❌ branch=%s, file=%s, err=%s (%.1fms)",
                        branch_name, rel_path, str(e), elapsed, exc_info=True)
            return None, None

    def add_file(
        self,
        branch_name: str,
        file_path: str,
        file_content: bytes = None,
        target_name: str = None
    ) -> Tuple[bool, str, str]:
        """向分支添加文件。
        
        Args:
            branch_name: 目标分支
            file_path: 源文件路径或文件名
            file_content: 文件内容（可选）
            target_name: 目标文件名（可选，默认使用原文件名；允许带'subdir/name.ext'写入子目录）
            
        Returns:
            (成功, 消息, 文件哈希)
        """
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            logger.error("[文件添加] 失败: 分支 '%s' 不存在", branch_name)
            return False, f"分支 '{branch_name}' 不存在", ""
        
        if target_name is None:
            target_name = Path(file_path).name
        
        target_path = branch_path / target_name
        source_size = len(file_content) if file_content else (Path(file_path).stat().st_size if Path(file_path).exists() else 0)
        
        logger.info("[文件添加] 请求: branch=%s, target=%s, size=%dKB", 
                    branch_name, target_name, source_size // 1024)
        
        try:
            if file_content is not None:
                with open(target_path, "wb") as f:
                    f.write(file_content)
                logger.info("[文件添加] 写入: %dKB 内容到 %s", source_size // 1024, target_path)
            else:
                source = Path(file_path)
                if source.exists():
                    shutil.copy2(source, target_path)
                    logger.info("[文件添加] 复制: %s -> %s", source, target_path)
                else:
                    logger.error("[文件添加] 失败: 源文件不存在 %s", file_path)
                    return False, f"源文件不存在: {file_path}", ""
            
            file_hash = self._compute_file_hash(target_path)
            file_size = target_path.stat().st_size
            self._add_version_record(branch_name, "commit", 
                f"添加文件: {target_name}", 1)
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件添加] 成功: %s (branch=%s, hash=%s..., size=%dKB, %.1fms)", 
                       target_name, branch_name, file_hash[:16], file_size // 1024, elapsed)
            return True, f"文件 '{target_name}' 已添加到分支 '{branch_name}'", file_hash
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件添加] 异常: %s -> %s (%.1fms)", file_path, str(e), elapsed)
            return False, f"添加文件失败: {str(e)}", ""
    
    def remove_file(self, branch_name: str, file_name: str) -> Tuple[bool, str]:
        """从分支删除文件。"""
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        target_path = branch_path / file_name

        logger.info("[文件删除] 开始操作: branch=%s, file=%s", branch_name, file_name)

        if not target_path.exists():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件删除] 文件不存在: branch=%s, file=%s (%.1fms)",
                           branch_name, file_name, elapsed)
            return False, f"文件 '{file_name}' 不存在于分支 '{branch_name}'"

        file_size = target_path.stat().st_size
        try:
            target_path.unlink()
            self._add_version_record(branch_name, "commit",
                f"删除文件: {file_name}", -1)
            # 同步DB一致性快照
            self._sync_snapshot_to_db(branch_name, change_note=f"rmfile:{file_name}")
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件删除] 成功: branch=%s, file=%s, size=%dKB (%.1fms)",
                       branch_name, file_name, file_size // 1024, elapsed)
            return True, f"文件 '{file_name}' 已从分支 '{branch_name}' 删除"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件删除] 异常: branch=%s, file=%s, err=%s (%.1fms)",
                        branch_name, file_name, str(e), elapsed, exc_info=True)
            return False, f"删除失败: {str(e)}"

    def rename_file(self, branch_name: str, old_path: str, new_name: str) -> Tuple[bool, str]:
        """重命名分支中的文件（支持同级目录改名或移动到同级子目录）。"""
        import time
        t0 = time.perf_counter()
        branch_path = self._get_branch_path(branch_name)
        old_target = branch_path / old_path

        logger.info("[文件重命名] 开始操作: branch=%s, old=%s -> new=%s",
                    branch_name, old_path, new_name)

        if not old_target.exists():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件重命名] 源文件不存在: branch=%s, old=%s (%.1fms)",
                           branch_name, old_path, elapsed)
            return False, f"源文件 '{old_path}' 不存在"

        if old_target.is_dir():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件重命名] 源是目录而非文件: branch=%s, old=%s (%.1fms)",
                           branch_name, old_path, elapsed)
            return False, f"'{old_path}' 是目录，暂不支持目录重命名"

        # 处理 new_name：可能仅文件名，也可能带子目录
        new_path = Path(new_name)
        if new_path.is_absolute():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件重命名] 新路径不能为绝对路径: branch=%s, new=%s (%.1fms)",
                           branch_name, new_name, elapsed)
            return False, "新路径不能使用绝对路径"

        new_target = branch_path / new_path
        # 必须保持在 branch_path 内（防穿越）
        try:
            new_target.resolve().relative_to(branch_path.resolve())
        except ValueError:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件重命名] 新路径越界: branch=%s, old=%s, new=%s (%.1fms)",
                           branch_name, old_path, new_name, elapsed)
            return False, "新路径不能跳出分支目录"

        if new_target.exists():
            elapsed = (time.perf_counter() - t0) * 1000
            logger.warning("[文件重命名] 目标已存在: branch=%s, target=%s (%.1fms)",
                           branch_name, str(new_path), elapsed)
            return False, f"目标文件 '{new_name}' 已存在，请先删除或换个名字"

        # 确保目标父目录存在
        new_target.parent.mkdir(parents=True, exist_ok=True)

        old_size = old_target.stat().st_size
        try:
            old_target.rename(new_target)
            self._add_version_record(branch_name, "commit",
                f"重命名文件: {old_path} -> {new_name}", -1)
            # 同步DB一致性快照（移动/重命名都需要）
            self._sync_snapshot_to_db(branch_name, change_note=f"mvfile:{old_path}->{new_name}")
            elapsed = (time.perf_counter() - t0) * 1000
            logger.info("[文件重命名] 成功: branch=%s, %s -> %s, size=%dKB (%.1fms)",
                       branch_name, old_path, new_name, old_size // 1024, elapsed)
            return True, f"已重命名: '{old_path}' → '{new_name}'"
        except Exception as e:
            elapsed = (time.perf_counter() - t0) * 1000
            logger.error("[文件重命名] 异常: branch=%s, old=%s, new=%s, err=%s (%.1fms)",
                        branch_name, old_path, new_name, str(e), elapsed, exc_info=True)
            return False, f"重命名失败: {str(e)}"
    
    def list_branch_files(self, branch_name: str) -> List[Dict]:
        """列出分支中的所有文件。"""
        branch_path = self._get_branch_path(branch_name)
        if not branch_path.exists():
            return []
        
        files = []
        for file_path in sorted(branch_path.rglob("*")):
            if file_path.is_file():
                relative_path = file_path.relative_to(branch_path)
                file_hash = self._compute_file_hash(file_path)
                is_valid = self._validate_file(file_path)
                file_size = file_path.stat().st_size
                
                files.append({
                    "path": str(relative_path),
                    "name": file_path.name,
                    "hash": file_hash,
                    "size": file_size,
                    "is_valid": is_valid,
                    "type": file_path.suffix.lower(),
                    "modified": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                })
        
        return files
    
    def detect_duplicates(self, branch_name: str) -> Dict:
        """检测分支中的重复文件。
        
        Returns:
            {
                "duplicates": [{"hash": ..., "files": [...]}],
                "total_files": int,
                "duplicate_groups": int,
                "wasted_space": int
            }
        """
        files = self.list_branch_files(branch_name)
        
        # 按哈希分组
        hash_groups: Dict[str, List[Dict]] = {}
        for f in files:
            h = f["hash"]
            if h not in hash_groups:
                hash_groups[h] = []
            hash_groups[h].append(f)
        
        # 找出重复的
        duplicates = []
        wasted_space = 0
        for h, group in hash_groups.items():
            if len(group) > 1:
                duplicates.append({
                    "hash": h,
                    "files": [f["path"] for f in group],
                    "count": len(group),
                    "size": group[0]["size"]
                })
                wasted_space += group[0]["size"] * (len(group) - 1)
        
        return {
            "duplicates": sorted(duplicates, key=lambda x: x["count"], reverse=True),
            "total_files": len(files),
            "duplicate_groups": len(duplicates),
            "wasted_space": wasted_space
        }
    
    def validate_branch(self, branch_name: str) -> Dict:
        """校验分支文件完整性。
        
        Returns:
            {
                "valid_files": int,
                "invalid_files": [{"path": ..., "error": ...}],
                "total_files": int,
                "issues": ["文件损坏: ...", ...]
            }
        """
        files = self.list_branch_files(branch_name)
        invalid_files = []
        valid_count = 0
        
        for f in files:
            if f["is_valid"]:
                valid_count += 1
            else:
                invalid_files.append({
                    "path": f["path"],
                    "error": f.get("error", "文件无法读取或已损坏")
                })
        
        return {
            "valid_files": valid_count,
            "invalid_files": invalid_files,
            "total_files": len(files),
            "issues": [f"文件损坏: {ifile['path']}" for ifile in invalid_files]
        }
    
    def compare_branches(
        self,
        branch_a: str,
        branch_b: str
    ) -> Dict:
        """比较两个分支的文件差异。
        
        Returns:
            {
                "only_in_a": [...],
                "only_in_b": [...],
                "modified": [...],
                "unchanged": [...],
                "summary": {...}
            }
        """
        files_a = {f["path"]: f for f in self.list_branch_files(branch_a)}
        files_b = {f["path"]: f for f in self.list_branch_files(branch_b)}
        
        paths_a = set(files_a.keys())
        paths_b = set(files_b.keys())
        
        only_in_a = [files_a[p] for p in paths_a - paths_b]
        only_in_b = [files_b[p] for p in paths_b - paths_a]
        
        common = paths_a & paths_b
        modified = []
        unchanged = []
        
        for p in common:
            if files_a[p]["hash"] != files_b[p]["hash"]:
                modified.append({
                    "path": p,
                    "hash_a": files_a[p]["hash"],
                    "hash_b": files_b[p]["hash"],
                    "size_a": files_a[p]["size"],
                    "size_b": files_b[p]["size"]
                })
            else:
                unchanged.append(files_a[p])
        
        return {
            "only_in_a": only_in_a,
            "only_in_b": only_in_b,
            "modified": modified,
            "unchanged": unchanged,
            "summary": {
                f"only_in_{branch_a}": len(only_in_a),
                f"only_in_{branch_b}": len(only_in_b),
                "modified": len(modified),
                "unchanged": len(unchanged)
            }
        }
    
    def merge_branches(
        self,
        source_branch: str,
        target_branch: str,
        strategy: str = "auto"
    ) -> Tuple[bool, Dict]:
        """合并分支。
        
        Args:
            source_branch: 源分支（被合并）
            target_branch: 目标分支（接收合并）
            strategy: 冲突解决策略
                - auto: 自动合并（无冲突直接合并，有冲突返回待处理）
                - keep_source: 遇到冲突保留源分支文件
                - keep_target: 遇到冲突保留目标分支文件
                - manual: 手动解决（返回冲突列表供外部处理）
                
        Returns:
            (成功, 结果字典)
        """
        result = {
            "merged_files": 0,
            "conflicts": [],
            "added_files": [],
            "updated_files": [],
            "errors": []
        }
        
        source_path = self._get_branch_path(source_branch)
        target_path = self._get_branch_path(target_branch)
        
        if not source_path.exists():
            return False, {"error": f"源分支 '{source_branch}' 不存在"}
        if not target_path.exists():
            return False, {"error": f"目标分支 '{target_branch}' 不存在"}
        
        source_files = {f["path"]: f for f in self.list_branch_files(source_branch)}
        target_files = {f["path"]: f for f in self.list_branch_files(target_branch)}
        
        # 处理源分支的每个文件
        for path, src_file in source_files.items():
            src_path = source_path / path
            tgt_path = target_path / path
            
            if path not in target_files:
                # 目标分支没有此文件 - 直接添加
                try:
                    tgt_path.parent.mkdir(parents=True, exist_ok=True)
                    shutil.copy2(src_path, tgt_path)
                    result["added_files"].append(path)
                    result["merged_files"] += 1
                except Exception as e:
                    result["errors"].append(f"添加文件 {path} 失败: {str(e)}")
            elif src_file["hash"] != target_files[path]["hash"]:
                # 文件存在但内容不同 - 冲突
                if strategy == "keep_source":
                    try:
                        shutil.copy2(src_path, tgt_path)
                        result["updated_files"].append(path)
                        result["merged_files"] += 1
                    except Exception as e:
                        result["errors"].append(f"更新文件 {path} 失败: {str(e)}")
                elif strategy == "keep_target":
                    # 保留目标分支，跳过
                    pass
                else:  # auto 或 manual
                    conflict = {
                        "path": path,
                        "source_hash": src_file["hash"],
                        "target_hash": target_files[path]["hash"],
                        "source_size": src_file["size"],
                        "target_size": target_files[path]["size"],
                        "resolution": "pending"
                    }
                    result["conflicts"].append(conflict)
        
        # 处理目标分支中有但源分支没有的文件（删除）
        for path in target_files:
            if path not in source_files:
                src_path = source_path / path
                if not src_path.exists():
                    tgt_path = target_path / path
                    try:
                        if tgt_path.exists():
                            # 根据策略决定是否删除
                            if strategy in ("auto", "keep_target"):
                                pass  # 保留目标分支的独有文件
                            elif strategy == "keep_source":
                                tgt_path.unlink()
                                result["removed_files"] = result.get("removed_files", []) + [path]
                    except Exception as e:
                        result["errors"].append(f"处理文件 {path} 失败: {str(e)}")
        
        # 记录版本
        self._add_version_record(target_branch, "merge",
            f"合并分支 '{source_branch}': 新增 {len(result['added_files'])} 个, "
            f"更新 {len(result['updated_files'])} 个, 冲突 {len(result['conflicts'])} 个",
            result["merged_files"])
        
        return len(result["errors"]) == 0, result
    
    def resolve_conflict(
        self,
        target_branch: str,
        file_path: str,
        resolution: str
    ) -> Tuple[bool, str]:
        """解决合并冲突。
        
        Args:
            target_branch: 目标分支
            file_path: 冲突文件路径
            resolution: 解决方案 (keep_source/keep_target)
        """
        branch_path = self._get_branch_path(target_branch)
        target_file = branch_path / file_path
        
        if resolution == "keep_target":
            # 保留当前目标分支的文件（什么都不做）
            return True, f"已保留目标分支的文件 '{file_path}'"
        elif resolution == "keep_source":
            # 需要从源分支获取最新文件
            # 这里简化处理：记录解决方案
            return True, f"已标记保留源分支版本的 '{file_path}'"
        else:
            return False, f"无效的解决方案: {resolution}"
    
    def get_branch_history(self, branch_name: str) -> List[Dict]:
        """获取分支的变更历史。"""
        meta = self._read_branch_metadata(branch_name)
        return meta.get("versions", [])
    
    def _compute_file_hash(self, file_path: Path) -> str:
        """计算文件的 SHA256 哈希。"""
        sha256 = hashlib.sha256()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(8192), b""):
                    sha256.update(chunk)
            return sha256.hexdigest()
        except Exception:
            return "ERROR"
    
    def _validate_file(self, file_path: Path) -> Tuple[bool, str]:
        """校验文件是否可读且未损坏。"""
        try:
            if not file_path.exists():
                return False, "文件不存在"
            if file_path.stat().st_size == 0:
                return False, "文件为空"
            
            # 尝试读取文件
            with open(file_path, "rb") as f:
                # 对于文本文件检查编码
                if file_path.suffix.lower() in (".csv", ".txt", ".json", ".md", ".py"):
                    content = f.read(1024)
                    try:
                        content.decode("utf-8")
                    except UnicodeDecodeError:
                        # 尝试其他编码
                        try:
                            content.decode("gbk")
                        except UnicodeDecodeError:
                            return False, "文件编码异常"
                else:
                    # 二进制文件只检查是否能读取
                    f.read(1)
            
            return True, ""
        except PermissionError:
            return False, "无读取权限"
        except Exception as e:
            return False, f"读取异常: {str(e)}"
    
    def _add_version_record(
        self,
        branch_name: str,
        change_type: str,
        summary: str,
        changed_files: int
    ):
        """添加版本记录。"""
        meta = self._read_branch_metadata(branch_name)
        versions = meta.get("versions", [])
        version_number = len(versions) + 1
        
        record = {
            "version": version_number,
            "change_type": change_type,
            "summary": summary,
            "changed_files": changed_files,
            "created_at": datetime.now().isoformat()
        }
        
        versions.append(record)
        meta["versions"] = versions
        self._update_branch_metadata(branch_name, {"versions": versions})


def get_branch_manager() -> BranchManager:
    """获取全局分支管理器实例。
    
    冷启动保护: 若 .branches/ 目录下尚无任何分支 meta 文件,
    自动创建默认 'main' 分支并激活, 避免 UI 中空分支下拉框导致 disabled。
    """
    global _branch_manager
    if "_branch_manager" not in globals():
        bm = BranchManager()
        existing = bm.list_branches()
        if not existing:
            import logging as _lgg
            _lgg.getLogger(__name__).info(
                "[分支管理] 冷启动: 尚无任何分支, 自动创建默认 main 分支并激活"
            )
            ok_create, _ = bm.create_branch("main", description="默认主分支 (冷启动自动创建)")
            if ok_create:
                bm.switch_branch("main")
        _branch_manager = bm
    return _branch_manager
