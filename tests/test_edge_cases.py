"""边界情况补充测试:覆盖 test_logging.py 未触达的分支。

运行: python tests/test_edge_cases.py

设计原则:
- 每个用例 try/except 捕获,明确标注 "预期外崩溃" 以便识别真实 Bug
- docx mock 文件动态生成,避免硬编码二进制
- 用例编号接 test_logging.py 的 9 号继续,从 10 开始
"""
from __future__ import annotations

import logging
import shutil
import sys
from pathlib import Path

import pandas as pd

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from src.log_config import setup_logging, get_logger
from src.data_loader import (
    parse_csv_filename,
    load_vehicle_csvs,
    load_durability_docx,
    load_durability_metadata,
    mark_invalid,
)
from src.metrics import (
    vehicle_overview,
    cell_voltage_consistency,
    power_summary,
    h2_system,
    fault_time_series,
    vehicle_speed_profile,
)
from src.plots import (
    fig_cell_voltage,
    fig_power_curve,
    fig_speed_hydrogen,
    fig_fault_bar,
    fig_durability_trend,
    fig_compare_overlay,
)
from src.report import _downsample, _flat_dict, build_report_html

setup_logging(level=logging.DEBUG)

TESTS_DIR = ROOT / "tests"
TMP_DIR = TESTS_DIR / "_tmp_edge"
TMP_DIR.mkdir(exist_ok=True)


def banner(title: str) -> None:
    print("\n" + "=" * 70)
    print(f"  {title}")
    print("=" * 70)


def expect_ok(name: str, fn):
    """执行 fn,正常返回则视为通过;若抛异常则打印 '预期外崩溃'。"""
    try:
        result = fn()
        print(f"  [OK] {name}")
        return result
    except Exception as e:
        print(f"  [预期外崩溃] {name}: {type(e).__name__}: {e}")
        return None


def main() -> None:
    # ============================================================
    # 一、log_config 模块
    # ============================================================
    banner("用例 10: log_config 重复调用 setup_logging 不应添加重复 handler")
    root = logging.getLogger()
    n_before = len(root.handlers)
    setup_logging()  # 第二次调用
    setup_logging()
    n_after = len(root.handlers)
    print(f"  → 调用前 handler 数={n_before}, 调用两次后={n_after} (应相等)")
    assert n_before == n_after, "重复调用导致 handler 重复添加!"

    banner("用例 11: setup_logging(log_file=...) 自动创建父目录并写文件")
    log_path = TMP_DIR / "subdir_not_exist" / "app.log"
    if log_path.parent.exists():
        shutil.rmtree(log_path.parent)
    # 重新初始化全局状态以触发文件 handler 路径
    import src.log_config as lc
    lc._configured = False
    # 清空已有 handler,模拟全新进程
    root2 = logging.getLogger()
    for h in list(root2.handlers):
        root2.removeHandler(h)
    setup_logging(level=logging.DEBUG, log_file=str(log_path))
    lg = get_logger("edge_test_file")
    lg.info("这是一条测试日志,验证文件写入")
    print(f"  → 日志文件存在: {log_path.exists()} (应为 True)")
    print(f"  → 父目录自动创建: {log_path.parent.exists()} (应为 True)")
    assert log_path.exists(), "日志文件未生成"

    banner("用例 12: get_logger 在 _configured=False 时自动初始化")
    import src.log_config as lc2
    lc2._configured = False
    root3 = logging.getLogger()
    for h in list(root3.handlers):
        root3.removeHandler(h)
    lg2 = get_logger("auto_init_test")
    print(f"  → _configured 状态: {lc2._configured} (应为 True)")
    assert lc2._configured, "get_logger 未自动初始化"

    # ============================================================
    # 二、data_loader 模块
    # ============================================================
    banner("用例 13: parse_csv_filename 含序号 (1) 应正确解析 seq=1")
    meta = parse_csv_filename("201480_202607071800_202607072359_CH0_20260807_225246 (1).csv")
    print(f"  → seq: {meta['seq'] if meta else 'None'} (应为 1)")
    assert meta and meta["seq"] == 1, "seq 未正确解析"

    banner("用例 14: parse_csv_filename 含非法日期应触发 ERROR 返回 None")
    # 月份 13 非法,to_datetime 应抛异常进入 except
    meta = parse_csv_filename("201480_202613321800_202607072359_CH0_20260807_225246.csv")
    print(f"  → 返回: {meta} (应为 None)")

    banner("用例 15: load_vehicle_csvs 传入空列表应返回空 DataFrame 不崩")
    df = expect_ok("空列表", lambda: load_vehicle_csvs([]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'} (应为 0)")

    banner("用例 16: mark_invalid 传入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: mark_invalid(pd.DataFrame()))

    banner("用例 17: mark_invalid 指定 columns 参数应只标记指定列")
    df_src = pd.DataFrame({
        "FC_MinCellVoltage": [800.0, 65535.0, -1.0],
        "FC_MaxCellVoltage": [805.0, 65535.0, 900.0],
        "OtherCol": [1, 2, 3],
    })
    flagged = mark_invalid(df_src, columns=["FC_MinCellVoltage"])
    has_min = "__FC_MinCellVoltage_invalid" in flagged.columns
    has_max = "__FC_MaxCellVoltage_invalid" in flagged.columns
    print(f"  → 标记 FC_MinCellVoltage: {has_min} (应 True)")
    print(f"  → 标记 FC_MaxCellVoltage: {has_max} (应 False,因未指定)")
    assert has_min and not has_max

    # ---------- docx 测试:动态构造 mock docx ----------
    banner("用例 18-22: docx 解析边界情况(动态构造 mock 文件)")

    def make_docx(path: Path, *, n_rows: int = 65, n_cols: int = 14,
                  with_header: bool = True, has_table: bool = True,
                  bad_value: bool = False):
        """生成 mock docx,模拟耐久测试表。"""
        from docx import Document
        doc = Document()
        if not has_table:
            doc.add_paragraph("无表格的文档")
            doc.save(str(path))
            return
        # 元数据 4 行(每行 5 格,键在第 0 格,值在第 3 格)
        meta_rows = [
            ["开始时间", "", "", "2026-07-07 18:00:00", ""],
            ["结束时间", "", "", "2026-07-07 23:59:59", ""],
            ["系统名称", "", "", "氢燃料电池系统", ""],
            ["电堆型号", "", "", "FC-Stack-V2", ""],
        ]
        # 列头行(行 4),用重复合并单元格模拟全角括号
        if n_cols == 14:
            header = ["目标功率(kW)", "湿度", "温度", "净输出功率(kW)", "电堆电流(A)",
                      "平均单体电压(V)", "离均差", "空压机功耗(kW)", "水泵功耗(kW)",
                      "冷却水入口温度(℃)", "冷却水出口温度(℃)", "HFR", "LFR", "电压方差"]
        else:
            header = [f"列{i}" for i in range(n_cols)]
        rows_data = meta_rows + [header] if with_header else meta_rows
        # 数据行(5..末)
        for i in range(n_rows - len(rows_data)):
            row_vals = []
            for j in range(n_cols):
                if bad_value and j == 3 and i == 2:
                    row_vals.append("BAD_VALUE")  # 触发数值化失败
                else:
                    row_vals.append(str(50 + i + j * 0.1))
            rows_data.append(row_vals)
        table = doc.add_table(rows=len(rows_data), cols=n_cols)
        for ri, row in enumerate(rows_data):
            for ci, val in enumerate(row):
                table.cell(ri, ci).text = str(val)
        doc.save(str(path))

    # 用例 18: 文件名不含"耐久X-Y"正则
    bad_name_docx = TMP_DIR / "随机名称.docx"
    make_docx(bad_name_docx)
    banner("用例 18: docx 文件名未匹配耐久区间正则,走 stem 兜底")
    df = expect_ok("stem 兜底", lambda: load_durability_docx([str(bad_name_docx)]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'}")

    # 用例 19: docx 无表格
    no_table_docx = TMP_DIR / "耐久0-5_无表格.docx"
    make_docx(no_table_docx, has_table=False)
    banner("用例 19: docx 无表格触发 WARNING 跳过")
    df = expect_ok("无表格", lambda: load_durability_docx([str(no_table_docx)]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'} (应为 0)")

    # 用例 20: 表行数不足 5
    short_docx = TMP_DIR / "耐久0-5_短表.docx"
    make_docx(short_docx, n_rows=3)
    banner("用例 20: docx 表行数不足 5 触发 ERROR 跳过")
    df = expect_ok("表行数不足", lambda: load_durability_docx([str(short_docx)]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'} (应为 0)")

    # 用例 21: 列数与预期不符(13 列)
    bad_col_docx = TMP_DIR / "耐久0-5_列数不符.docx"
    make_docx(bad_col_docx, n_cols=13)
    banner("用例 21: docx 列数与预期不符触发 WARNING 用标准列名兜底")
    df = expect_ok("列数不符", lambda: load_durability_docx([str(bad_col_docx)]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'}")

    # 用例 22: 列值含非数字字符串触发数值化失败 WARNING
    bad_val_docx = TMP_DIR / "耐久0-5_含坏值.docx"
    make_docx(bad_val_docx, bad_value=True)
    banner("用例 22: docx 列值含非数字触发数值化失败 WARNING")
    df = expect_ok("含坏值", lambda: load_durability_docx([str(bad_val_docx)]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'}")

    # 用例 23: docx 文件不存在
    banner("用例 23: docx 文件不存在触发 ERROR")
    df = expect_ok("文件不存在", lambda: load_durability_docx([str(TMP_DIR / "no_such.docx")]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'} (应为 0)")

    # 用例 24: load_durability_metadata 文件不存在
    banner("用例 24: load_durability_metadata 文件不存在触发 ERROR 返回空")
    df = expect_ok("元数据文件不存在", lambda: load_durability_metadata([str(TMP_DIR / "no_such.docx")]))
    print(f"  → 返回行数: {len(df) if df is not None else 'N/A'} (应为 0)")

    # ============================================================
    # 三、metrics 模块
    # ============================================================
    banner("用例 25: vehicle_overview 传入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: vehicle_overview(pd.DataFrame()))

    banner("用例 26: cell_voltage_consistency 缺 FC_MinVoltageChannel 列应跳过最弱通道")
    df_no_ch = pd.DataFrame({
        "FC_MinCellVoltage": [800.0, 801.0],
        "FC_MaxCellVoltage": [805.0, 806.0],
        "FC_AvgCellVoltage": [802.0, 803.0],
    })
    res = expect_ok("缺最弱通道列", lambda: cell_voltage_consistency(df_no_ch))
    print(f"  → 是否含最弱通道Top5: {'最弱通道Top5' in res if res else 'N/A'} (应 False)")

    banner("用例 27: cell_voltage_consistency 全为异常值触发 WARNING 无数据")
    df_all_bad = pd.DataFrame({
        "FC_MinCellVoltage": [65535, -1, 999],
        "FC_MaxCellVoltage": [65535, -1, 999],
        "FC_AvgCellVoltage": [65535, -1, 999],
    })
    res = expect_ok("全异常值", lambda: cell_voltage_consistency(df_all_bad))
    print(f"  → 产出键数: {len(res) if res else 'N/A'} (应 0)")

    banner("用例 28: power_summary 全为异常值应返回空 dict")
    df_pw_bad = pd.DataFrame({"FC_NetPwrOut": [65535, -1], "FC_CurrOut": [65535, -1]})
    res = expect_ok("功率全异常", lambda: power_summary(df_pw_bad))
    print(f"  → 产出键数: {len(res) if res else 'N/A'} (应 0)")

    banner("用例 29: h2_system 传入空 DataFrame 应返回空 dict")
    expect_ok("空 DataFrame", lambda: h2_system(pd.DataFrame()))

    banner("用例 30: fault_time_series 无 FC_ErrorCode 列触发 WARNING 返回空")
    df_no_ec = pd.DataFrame({"Timestamp": pd.to_datetime(["2026-07-07 18:00:00"])})
    res = expect_ok("无 ErrorCode 列", lambda: fault_time_series(df_no_ec))
    print(f"  → 返回行数: {len(res) if res is not None else 'N/A'} (应 0)")

    banner("用例 31: fault_time_series 缺 FC_SysFltRnk 列(疑似 Bug,会 KeyError)")
    df_has_ec_no_rank = pd.DataFrame({
        "Timestamp": pd.to_datetime(["2026-07-07 18:00:00", "2026-07-07 18:01:00"]),
        "FC_ErrorCode": [0.0, 100.0],  # 第二条触发故障
        # 故意不提供 FC_SysFltRnk 列
    })
    res = expect_ok("缺 FC_SysFltRnk 列", lambda: fault_time_series(df_has_ec_no_rank))
    print(f"  → 返回行数: {len(res) if res is not None else 'N/A'}")

    banner("用例 32: vehicle_speed_profile 无车速/氢耗/里程列应只返回 Timestamp")
    df_no_speed = pd.DataFrame({"Timestamp": pd.to_datetime(["2026-07-07 18:00:00"])})
    res = expect_ok("无车速列", lambda: vehicle_speed_profile(df_no_speed))
    print(f"  → 返回列: {list(res.columns) if res is not None else 'N/A'} (应只含 Timestamp)")

    # ============================================================
    # 四、plots 模块
    # ============================================================
    banner("用例 33: fig_cell_voltage 输入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: fig_cell_voltage(pd.DataFrame()))

    banner("用例 34: fig_cell_voltage 输入无指定列应渲染空图")
    df_no_cols = pd.DataFrame({"Timestamp": pd.to_datetime(["2026-07-07 18:00:00"])})
    fig = expect_ok("无指定列", lambda: fig_cell_voltage(df_no_cols))
    print(f"  → 图中 trace 数: {len(fig.data) if fig else 'N/A'} (应 0)")

    banner("用例 35: fig_power_curve 输入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: fig_power_curve(pd.DataFrame()))

    banner("用例 36: fig_speed_hydrogen 输入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: fig_speed_hydrogen(pd.DataFrame()))

    banner("用例 37: fig_durability_trend 输入空 DataFrame 应不崩")
    expect_ok("空 DataFrame", lambda: fig_durability_trend(pd.DataFrame()))

    banner("用例 38: fig_compare_overlay df_a 为空应不崩")
    df_a_empty = pd.DataFrame()
    df_b_ok = pd.DataFrame({
        "Timestamp": pd.to_datetime(["2026-07-07 18:00:00"]),
        "FC_NetPwrOut": [10.0],
    })
    fig = expect_ok("df_a 空", lambda: fig_compare_overlay(df_a_empty, df_b_ok, "FC_NetPwrOut"))
    print(f"  → 图中 trace 数: {len(fig.data) if fig else 'N/A'} (应 1)")

    banner("用例 39: fig_compare_overlay 列在 df_a 不在 df_b")
    df_a_has = pd.DataFrame({
        "Timestamp": pd.to_datetime(["2026-07-07 18:00:00"]),
        "FC_NetPwrOut": [10.0],
    })
    df_b_no = pd.DataFrame({"Timestamp": pd.to_datetime(["2026-07-07 18:00:00"])})
    fig = expect_ok("列只在 df_a", lambda: fig_compare_overlay(df_a_has, df_b_no, "FC_NetPwrOut"))
    print(f"  → 图中 trace 数: {len(fig.data) if fig else 'N/A'} (应 1)")

    banner("用例 40: fig_cell_voltage 列含字符串(疑似 Bug,会 TypeError)")
    # plots.py 第 27 行 sub[col] > 0 直接比较,未用 _safe_num
    df_str = pd.DataFrame({
        "Timestamp": pd.to_datetime(["2026-07-07 18:00:00", "2026-07-07 18:01:00"]),
        "FC_MinCellVoltage": ["NOT_A_NUMBER", "800.0"],
    })
    fig = expect_ok("字符串列", lambda: fig_cell_voltage(df_str))

    # ============================================================
    # 五、report 模块
    # ============================================================
    banner("用例 41: _downsample 行数等于 max_points 应原样返回")
    df_eq = pd.DataFrame({"x": range(1000)})
    out = expect_ok("行数等于", lambda: _downsample(df_eq, max_points=1000))
    print(f"  → 返回行数: {len(out) if out is not None else 'N/A'} (应 1000)")

    banner("用例 42: _downsample 行数小于 max_points 应原样返回")
    df_lt = pd.DataFrame({"x": range(100)})
    out = expect_ok("行数小于", lambda: _downsample(df_lt, max_points=1000))
    print(f"  → 返回行数: {len(out) if out is not None else 'N/A'} (应 100)")

    banner("用例 43: _flat_dict 传入非 dict 应返回空 dict")
    out = expect_ok("非 dict", lambda: _flat_dict("not a dict"))
    print(f"  → 返回: {out} (应 {{}})")

    banner("用例 44: _flat_dict 传入空 dict 应返回空 dict")
    out = expect_ok("空 dict", lambda: _flat_dict({}))
    print(f"  → 返回: {out} (应 {{}})")

    banner("用例 45: build_report_html 输入空 DataFrame 应能生成报告")
    df_empty = pd.DataFrame({"Timestamp": pd.to_datetime([])})
    html = expect_ok("空 DataFrame 报告", lambda: build_report_html(
        "TEST001", df_empty, {}, {}, {}, {}))
    print(f"  → HTML 长度: {len(html) if html else 'N/A'}")

    # 清理临时目录(先关闭文件 handler 句柄,避免 WinError 32)
    try:
        root_end = logging.getLogger()
        for h in list(root_end.handlers):
            if isinstance(h, logging.FileHandler):
                h.close()
                root_end.removeHandler(h)
        shutil.rmtree(TMP_DIR, ignore_errors=True)
        if TMP_DIR.exists():
            print(f"\n[清理] 临时目录仍存在: {TMP_DIR}(可能有句柄未释放)")
        else:
            print("\n[清理] 临时目录已删除")
    except Exception as e:
        print(f"\n[清理] 删除失败: {e}")

    print("\n" + "=" * 70)
    print("  边界情况测试已全部执行,请检查上方 [预期外崩溃] 标记")
    print("=" * 70 + "\n")


if __name__ == "__main__":
    main()
