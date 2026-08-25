"""报告生成模块:HTML 报告模板,可浏览器打印为 PDF;一键整车报告生成。"""
from __future__ import annotations

import base64
import io
import logging
from pathlib import Path
from datetime import datetime

import pandas as pd
import plotly.graph_objects as go

from src.log_config import get_logger
from src.plots import (
    fig_cell_voltage,
    fig_fault_bar,
    fig_power_curve,
    fig_speed_hydrogen,
)

logger = get_logger(__name__)


# =====================================================================
# 辅助函数 (放在所有业务函数之前, 避免同模块内部出现自引用 import)
# =====================================================================

def _fig_to_html(fig: go.Figure, full_html: bool = False) -> str:
    """Plotly 图 → 内嵌 HTML (关闭工具栏减少 DOM 体积)。"""
    return fig.to_html(
        include_plotlyjs="cdn",
        full_html=full_html,
        div_id=None,
        config={"displayModeBar": False, "responsive": True},
    )


def _downsample(df: pd.DataFrame, max_points: int = 1000) -> pd.DataFrame:
    """大数据量降采样到 max_points 条(等间隔抽样)。"""
    if len(df) <= max_points:
        return df
    step = max(1, len(df) // max_points)
    out = df.iloc[::step].reset_index(drop=True)
    logger.info("降采样: %d → %d (step=%d)", len(df), len(out), step)
    return out


def _flat_dict(d: dict, prefix: str = "") -> dict:
    """扁平化嵌套字典,键含路径,便于行表展示。"""
    out: dict[str, str] = {}
    if not isinstance(d, dict):
        return out
    for k, v in d.items():
        key = f"{prefix}.{k}" if prefix else str(k)
        if isinstance(v, dict):
            out.update(_flat_dict(v, key))
        else:
            out[key] = str(v)
    return out


# =====================================================================
# 业务入口:一键整车报告
# =====================================================================

def generate_vehicle_report(
    vehicle: str,
    df: pd.DataFrame,
    *,
    fmt: str = "html",
) -> tuple[str, bytes | str]:
    """一键生成整车报告。

    Args:
        vehicle: 车辆编号
        df: 整车 DataFrame (load_vehicle_csvs 输出格式)
        fmt:  "html" | "docx" (docx 会生成最小可行 Word,图片用 Plotly 静态 PNG 插入)

    Returns:
        (file_name, content)
        - fmt=html: content 是 str,直接 download_button 使用 encode('utf-8')
        - fmt=docx: content 是 bytes,直接给 download_button
    """
    from src.metrics import (
        vehicle_overview,
        cell_voltage_consistency,
        power_summary,
        h2_system,
    )

    logger.info("[报告] generate_vehicle_report 开始: 车辆=%s 行数=%d fmt=%s",
                vehicle, 0 if df is None else len(df), fmt)
    if df is None or len(df) == 0:
        raise ValueError("数据为空,无法生成报告")

    t0 = datetime.now()

    # ---- 计算 4 组指标 (懒导入,避免 report 模块顶层依赖 metrics) ----
    overview = vehicle_overview(df)
    cell_consist = cell_voltage_consistency(df)
    power = power_summary(df)
    h2 = h2_system(df)
    calc_ms = int((datetime.now() - t0).total_seconds() * 1000)
    logger.info("[报告] 指标计算完成: overview_keys=%d cell=%d power=%d h2=%d 耗时=%dms",
                len(overview), len(cell_consist), len(power), len(h2), calc_ms)

    # ---- 组装 HTML (统一入口, docx 会再从 HTML 生成或另起 Word pipeline) ----
    html_content = build_report_html(vehicle, df, overview, cell_consist, power, h2)
    base_name = f"vehicle_{vehicle}_{datetime.now():%Y%m%d_%H%M%S}"

    if fmt.lower() == "html":
        logger.info("[报告] HTML 生成完成: len=%d 字符", len(html_content))
        return f"{base_name}.html", html_content

    if fmt.lower() == "docx":
        docx_bytes = _html_to_minimal_docx(
            html_content, vehicle, df, overview, cell_consist, power, h2,
        )
        logger.info("[报告] DOCX 生成完成: size=%d bytes", len(docx_bytes))
        return f"{base_name}.docx", docx_bytes

    raise ValueError(f"不支持的报告格式: {fmt}, 请使用 html / docx")


def _fig_to_png_bytes(fig: go.Figure, width: int = 1000, height: int = 500) -> bytes:
    """Plotly 图 → PNG bytes。优先用 kaleido;失败时返回空 bytes (上层会显示占位文字)。"""
    try:
        return fig.to_image(
            format="png",
            width=width,
            height=height,
            engine="kaleido",
        )
    except Exception as e:
        logger.warning("[报告] kaleido 导出 PNG 失败(%s),跳过嵌入静态图", e)
        return b""


def _html_to_minimal_docx(
    html: str,
    vehicle: str,
    df: pd.DataFrame,
    overview: dict,
    cell_consist: dict,
    power: dict,
    h2: dict,
) -> bytes:
    """生成最小可行 Word 报告 (KPI 表 + 指标明细 + 4 张图)。

    python-docx 不需要浏览器,Streamlit Cloud 友好。
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- 标题 ----
    title = doc.add_heading("燃料电池整车测试报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_p = doc.add_paragraph()
    meta_p.add_run("车辆编号: ").bold = True
    meta_p.add_run(f"{vehicle}    ")
    meta_p.add_run("生成时间: ").bold = True
    meta_p.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))

    # ---- 一、KPI 概览 (2x4 表格) ----
    doc.add_heading("一、关键指标概览", level=1)
    kpi_keys = [
        ("运行时长(h)", "运行时长(h)"),
        ("行驶里程(km)", "行驶里程(km)"),
        ("平均车速(km/h)", "平均车速(km/h)"),
        ("百公里氢耗均值(kg)", "百公里氢耗均值(kg)"),
        ("启动次数", "启动次数"),
        ("故障码种类", "故障码种类"),
        ("采样点数", "采样点数"),
        ("里程末值(km)", "里程末值(km)"),
    ]
    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Light Grid Accent 1"
    for i, (label, key) in enumerate(kpi_keys):
        r, c = divmod(i, 4)
        cell = tbl.rows[r].cells[c]
        cell.text = ""
        p_lab = cell.paragraphs[0]
        run_lab = p_lab.add_run(label + "\n")
        run_lab.font.size = Pt(9)
        run_lab.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run_val = p_lab.add_run(str(overview.get(key, "-")))
        run_val.bold = True
        run_val.font.size = Pt(14)
        run_val.font.color.rgb = RGBColor(0x1F, 0x77, 0xB4)

    # ---- 二、详细指标 (直接调用同模块辅助函数,不再自引用 import) ----
    doc.add_heading("二、详细指标", level=1)
    detail: dict[str, str] = {}
    detail.update(_flat_dict({"单片一致性": cell_consist}))
    detail.update(_flat_dict({"功率与效率": power}))
    detail.update(_flat_dict({"氢系统": h2}))
    detail_table = doc.add_table(rows=len(detail) + 1, cols=2)
    detail_table.style = "Light Grid Accent 1"
    hdr_cells = detail_table.rows[0].cells
    hdr_cells[0].text = "指标"
    hdr_cells[1].text = "数值"
    for idx, (k, v) in enumerate(detail.items(), 1):
        row = detail_table.rows[idx].cells
        row[0].text = k
        row[1].text = v

    # ---- 三、图表 (直接调用模块级名字, 不做二次 import) ----
    doc.add_heading("三、图表", level=1)
    df_s = _downsample(df, max_points=500)

    chart_specs = [
        ("3.1 单片电压一致性",        lambda: fig_cell_voltage(df_s)),
        ("3.2 功率与电流",            lambda: fig_power_curve(df_s)),
        ("3.3 车速与瞬时氢耗",        lambda: fig_speed_hydrogen(df_s)),
        ("3.4 故障码分布",            lambda: fig_fault_bar(overview.get("故障码Top10", {}))),
    ]
    inserted_any = False
    for ctitle, fig_fn in chart_specs:
        doc.add_heading(ctitle, level=2)
        try:
            fig = fig_fn()
        except Exception as e:
            logger.warning("[报告-DOCX] %s 图表 fig 生成失败(不中断): %s", ctitle, e)
            doc.add_paragraph(f"(无法生成图表: {type(e).__name__}: {str(e)[:80]})")
            continue
        png = _fig_to_png_bytes(fig)
        if png:
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(6.2))
                last = doc.paragraphs[-1]
                last.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted_any = True
            except Exception as e:
                logger.warning("[报告-DOCX] %s 插入图片失败: %s", ctitle, e)
                doc.add_paragraph(f"(图表渲染失败: {str(e)[:50]})")
        else:
            doc.add_paragraph("(静态图不可用: 缺 kaleido,建议配合 HTML 版 Ctrl+P 打印为高清 PDF)")

    if not inserted_any:
        p_note = doc.add_paragraph()
        run_note = p_note.add_run(
            "提示: 本地安装 pip install kaleido 后可以在 Word 中嵌入静态图; "
            "或直接下载 HTML 版本,浏览器打开后用 Ctrl+P 打印为高清 PDF。"
        )
        run_note.font.size = Pt(9)
        run_note.font.color.rgb = RGBColor(0x88, 0x55, 0x00)

    # ---- 四、结论 ----
    doc.add_heading("四、结论", level=1)
    dur = overview.get("运行时长(h)", 0)
    km = overview.get("行驶里程(km)", 0)
    h2c = overview.get("百公里氢耗均值(kg)", 0)
    fault_cnt = overview.get("故障码种类", 0)
    doc.add_paragraph(
        f"本次测试车辆 {vehicle},运行 {dur} 小时,行驶 {km} km,"
        f"百公里氢耗 {h2c} kg。"
        f"故障码种类 {fault_cnt} 种。"
        "建议结合单片电压一致性与故障时间轴,关注电压最低单体通道及高频故障码。"
    )

    # ---- 页脚 ----
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "本报告由『氢质氢离 · 设备测试数据分析与自动报告助手』自动生成。"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


HTML_TEMPLATE = """<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>燃料电池测试报告 - {vehicle}</title>
<style>
* {{ box-sizing: border-box; }}
body {{ font-family: "Microsoft YaHei", "PingFang SC", sans-serif; margin: 40px; color: #222; }}
h1 {{ font-size: 26px; border-bottom: 3px solid #1f77b4; padding-bottom: 10px; }}
h2 {{ font-size: 18px; color: #1f77b4; margin-top: 30px; border-left: 4px solid #1f77b4; padding-left: 8px; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
td, th {{ border: 1px solid #ddd; padding: 6px 10px; text-align: left; font-size: 13px; }}
th {{ background: #f5f5f5; }}
.kpi-grid {{ display: grid; grid-template-columns: repeat(4,1fr); gap: 10px; margin: 15px 0; }}
.kpi {{ background: #f9f9f9; padding: 12px; border-radius: 6px; border-left: 4px solid #1f77b4; }}
.kpi .label {{ font-size: 12px; color: #666; }}
.kpi .value {{ font-size: 22px; font-weight: bold; color: #1f77b4; }}
.chart {{ margin: 15px 0; }}
.note {{ background: #fff8e1; border-left: 4px solid #ff9800; padding: 10px 14px; margin: 10px 0; font-size: 13px; color: #6a4a00; border-radius: 4px; }}
.footer {{ margin-top: 40px; font-size: 12px; color: #888; border-top: 1px solid #ddd; padding-top: 10px; }}
@media print {{ body {{ margin: 15px; }} .no-print {{ display: none; }} }}
</style></head>
<body>
<h1>燃料电池整车测试报告</h1>
<p>车辆编号: <b>{vehicle}</b>  |  生成时间: {gen_time}</p>

<h2>一、关键指标概览</h2>
<div class="kpi-grid">{kpi_cards}</div>

<h2>二、详细指标</h2>
<table>{detail_rows}</table>

<h2>三、单片电压一致性</h2>
<div class="note">
  <b>⚠ 数据单位说明:</b> 经样本数据扫描,本数据集 <code>FC_MinCellVoltage / FC_MaxCellVoltage / FC_AvgCellVoltage</code> 字段数值集中在 100~1000 区间(典型值约 600~900),<b>单位疑似 mV 而非 V</b>(燃料电池单片电压物理范围 0.6~0.9 V,对应 600~900 mV)。<br>
  阅读本章节统计值时请按 mV 理解:例如 mean=739.57 应理解为 <b>0.74 V</b>。<br>
  压差 <code>cell_diff</code> 同理:例如 max=49.00 应理解为 <b>0.049 V (49 mV)</b>,属正常单片压差范围。
</div>
<div class="chart">{fig_cell}</div>

<h2>四、功率与电流</h2>
<div class="chart">{fig_power}</div>

<h2>五、车速与瞬时氢耗</h2>
<div class="chart">{fig_speed}</div>

<h2>六、故障码分布</h2>
<div class="chart">{fig_fault}</div>

<h2>七、结论</h2>
<p>{conclusion}</p>

<div class="footer">本报告由『氢质氢离 · 设备测试数据分析与自动报告助手』自动生成。</div>
</body></html>
"""


def build_report_html(
    vehicle: str,
    df: pd.DataFrame,
    overview: dict,
    cell_consist: dict,
    power: dict,
    h2: dict,
) -> str:
    """组装完整 HTML 报告。"""
    logger.info("开始组装 HTML 报告: 车辆=%s 行数=%d", vehicle, len(df))

    # KPI 卡片
    kpi_keys = [
        ("运行时长(h)", "运行时长(h)"),
        ("行驶里程(km)", "行驶里程(km)"),
        ("平均车速(km/h)", "平均车速(km/h)"),
        ("百公里氢耗均值(kg)", "百公里氢耗均值(kg)"),
        ("启动次数", "启动次数"),
        ("故障码种类", "故障码种类"),
        ("采样点数", "采样点数"),
        ("里程末值(km)", "里程末值(km)"),
    ]
    cards = []
    for label, key in kpi_keys:
        v = overview.get(key, "-")
        cards.append(
            f'<div class="kpi"><div class="label">{label}</div>'
            f'<div class="value">{v}</div></div>'
        )

    # 详细表
    detail: dict[str, str] = {}
    detail.update(_flat_dict({"单片一致性": cell_consist}))
    detail.update(_flat_dict({"功率与效率": power}))
    detail.update(_flat_dict({"氢系统": h2}))
    rows = "".join(
        f"<tr><td>{k}</td><td>{v}</td></tr>" for k, v in detail.items()
    )

    # 图表(降采样避免 HTML 体积过大,防御缺列异常)
    df_s = _downsample(df, max_points=1000)
    _safe_plots = [
        ("fig_cell",   fig_cell_voltage,      df_s,
         "三、单片电压一致性",        "缺少电压列或 Timestamp 列,无法绘制"),
        ("fig_power",  fig_power_curve,      df_s,
         "四、功率与电流",            "缺少功率相关列或 Timestamp 列,无法绘制"),
        ("fig_speed",  fig_speed_hydrogen,   df_s,
         "五、车速与瞬时氢耗",        "缺少车速/氢耗列或 Timestamp 列,无法绘制"),
    ]
    figs: dict[str, str] = {}
    for key, fn, arg, ch_title, ch_note in _safe_plots:
        try:
            figs[key] = _fig_to_html(fn(arg))
            logger.info("[报告] 图表 %s 生成成功", ch_title)
        except Exception as e:
            logger.warning("[报告] 图表 %s 生成失败(不中断报告): %s", ch_title, e)
            # 替换为占位 HTML (note样式,与整体主题一致)
            figs[key] = (
                f'<div class="note"><b>图表占位 · {ch_title}</b><br>'
                f'{ch_note}<br>'
                f'<span style="font-size:11px;color:#999">'
                f'Error: {type(e).__name__}: {str(e)[:120]}</span></div>'
            )
    # fig_fault_bar 函数已经内置了空 fault_top 的友好处理,但继续加 try 防未知异常
    try:
        figs["fig_fault"] = _fig_to_html(fig_fault_bar(overview.get("故障码Top10", {})))
    except Exception as e:
        logger.warning("[报告] 图表 六、故障码分布 生成失败(不中断): %s", e)
        figs["fig_fault"] = (
            f'<div class="note"><b>图表占位 · 六、故障码分布</b><br>'
            f'Error: {type(e).__name__}: {str(e)[:120]}</div>'
        )
    fig_cell = figs["fig_cell"]
    fig_power = figs["fig_power"]
    fig_speed = figs["fig_speed"]
    fig_fault = figs["fig_fault"]

    # 结论(自动)
    dur = overview.get("运行时长(h)", 0)
    km = overview.get("行驶里程(km)", 0)
    h2c = overview.get("百公里氢耗均值(kg)", 0)
    fault_cnt = overview.get("故障码种类", 0)
    conclusion = (
        f"本次测试车辆 {vehicle},运行 {dur} 小时,行驶 {km} km,"
        f"百公里氢耗 {h2c} kg。"
        f"故障码种类 {fault_cnt} 种。"
        "建议结合单片电压一致性与故障时间轴,关注电压最低单体通道及高频故障码。"
    )

    return HTML_TEMPLATE.format(
        vehicle=vehicle,
        gen_time=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        kpi_cards="".join(cards),
        detail_rows=rows,
        fig_cell=fig_cell,
        fig_power=fig_power,
        fig_speed=fig_speed,
        fig_fault=fig_fault,
        conclusion=conclusion,
    )
