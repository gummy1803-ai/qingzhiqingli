"""数据加载模块:CSV 批量拼接 + docx 耐久解析 + 文件名解析。"""
from __future__ import annotations

import logging
import re
import time
from pathlib import Path
from typing import Iterable

import pandas as pd

from src.log_config import get_logger

logger = get_logger(__name__)


# 文件名形如:201480_202607071800_202607072359_CH0_20260807_225246 (1).csv
# 含义:车辆编号_开始时间_结束时间_通道_导出时间_时分秒 (序号)
CSV_NAME_RE = re.compile(
    r"^(?P<vehicle>\d+)"
    r"_(?P<start>\d{12})"
    r"_(?P<end>\d{12})"
    r"_CH(?P<channel>\d+)"
    r"_(?P<export>\d{8}_\d{6})"
    r"(?:\s*\((?P<seq>\d+)\))?"
    r"\.csv$"
)


def parse_csv_filename(name: str) -> dict | None:
    """解析 CSV 文件名,返回车辆/起止时间/通道等元数据。"""
    logger.debug("解析 CSV 文件名: %s", name)
    m = CSV_NAME_RE.match(name)
    if not m:
        logger.warning("CSV 文件名不匹配规则,跳过: %s", name)
        return None
    g = m.groupdict()
    try:
        result = {
            "vehicle": g["vehicle"],
            "start_ts": pd.to_datetime(g["start"], format="%Y%m%d%H%M%S"),
            "end_ts": pd.to_datetime(g["end"], format="%Y%m%d%H%M%S"),
            "channel": int(g["channel"]),
            "export_ts": pd.to_datetime(g["export"], format="%Y%m%d_%H%M%S"),
            "seq": int(g["seq"]) if g["seq"] else 0,
        }
        logger.debug(
            "解析成功: vehicle=%s start=%s end=%s channel=%d seq=%d",
            result["vehicle"], result["start_ts"], result["end_ts"],
            result["channel"], result["seq"],
        )
        return result
    except Exception as e:
        logger.error("解析 CSV 文件名失败: %s | 错误: %s", name, e, exc_info=True)
        return None


def load_vehicle_csvs(file_paths: Iterable[str]) -> pd.DataFrame:
    """批量读取同一车辆 CSV 分片,按时间排序、去重。

    - 自动解析文件名元数据
    - Timestamp 列转 datetime
    - 按 Timestamp 排序、去重(保留首条)
    - 异常值标记(65535、负值),但不删除,由上层决定如何处理
    """
    paths = list(file_paths)
    logger.info("=== CSV 批量加载开始,共 %d 个文件 ===", len(paths))
    t_start = time.perf_counter()

    dfs: list[pd.DataFrame] = []
    meta_rows: list[dict] = []
    skip_count = 0
    ok_count = 0

    for i, fp in enumerate(paths):
        p = Path(fp)
        logger.info("[%d/%d] 处理: %s", i + 1, len(paths), p.name)
        meta = parse_csv_filename(p.name)
        if meta is None:
            logger.warning("  → 文件名未匹配,跳过元数据")
        try:
            df = pd.read_csv(p)
            logger.debug("  → 读取成功 shape=%s cols=%d", df.shape, len(df.columns))
        except Exception as e:
            logger.error("  → 读取失败: %s", e, exc_info=True)
            skip_count += 1
            continue
        if "Timestamp" not in df.columns:
            logger.warning("  → 缺少 Timestamp 列,跳过。实际列: %s",
                           list(df.columns)[:5])
            skip_count += 1
            continue
        # 转 datetime
        df["Timestamp"] = pd.to_datetime(df["Timestamp"], errors="coerce")
        n_na = int(df["Timestamp"].isna().sum())
        if n_na:
            logger.warning("  → Timestamp 解析有 %d 个 NaT,将剔除", n_na)
        df = df.dropna(subset=["Timestamp"]).sort_values("Timestamp")
        if meta:
            df["__vehicle"] = meta["vehicle"]
            df["__channel"] = meta["channel"]
            meta_rows.append(meta)
        logger.debug("  → 入库 %d 行,时间范围 %s ~ %s",
                     len(df), df["Timestamp"].iloc[0], df["Timestamp"].iloc[-1])
        dfs.append(df)
        ok_count += 1

    if not dfs:
        logger.warning("=== CSV 加载结束,无可用数据(全部跳过)===")
        return pd.DataFrame()

    logger.info("拼接 %d 个 DataFrame ...", len(dfs))
    merged = pd.concat(dfs, ignore_index=True)
    logger.info("拼接完成,共 %d 行", len(merged))

    # 同一 Timestamp 去重
    before_dedup = len(merged)
    merged = merged.drop_duplicates(subset=["Timestamp"], keep="first")
    after_dedup = len(merged)
    if before_dedup != after_dedup:
        logger.info("去重: %d → %d (剔除 %d 条重复)",
                    before_dedup, after_dedup, before_dedup - after_dedup)

    merged = merged.sort_values("Timestamp").reset_index(drop=True)
    elapsed = time.perf_counter() - t_start
    logger.info(
        "=== CSV 加载结束:成功 %d / 跳过 %d / 最终 %d 行 / 耗时 %.2fs ===",
        ok_count, skip_count, len(merged), elapsed,
    )
    return merged


# ---------- 异常值处理 ----------

# 标记类阈值:常见无效哨兵值
INVALID_SENTINELS = {65535, -1, 999, 99}


def mark_invalid(df: pd.DataFrame, columns: list[str] | None = None) -> pd.DataFrame:
    """在指定列追加 <col>__invalid 布尔列,标记哨兵值/负值。

    不修改原值,只打标,上层可决定 replace/丢弃。
    """
    logger.info("异常值打标: 输入 %d 行", len(df))
    cols = columns or [
        "FC_MinCellVoltage",
        "FC_MaxCellVoltage",
        "FC_AvgCellVoltage",
        "FC_CurrOut",
        "FC_VoltOut",
        "FC_NetPwrOut",
        "FC_VehicleSpd",
        "FC_VehicleKM",
        "FC_HydCmPerHundred",
        "FC_HydCmInstts",
    ]
    out = df.copy()
    flagged_cols = 0
    total_flagged = 0
    for c in cols:
        if c not in out.columns:
            continue
        # 先尝试数值化;非数值列只比较哨兵值,跳过 < 0 比较
        try:
            num = pd.to_numeric(out[c], errors="coerce")
            lt_zero = num < 0
        except Exception as e:
            logger.debug("  列 %s 无法数值化,仅比较哨兵值: %s", c, e)
            lt_zero = pd.Series([False] * len(out))
        flag = out[c].isin(INVALID_SENTINELS) | lt_zero
        n = int(flag.sum())
        if n:
            logger.warning("  列 %s 标记异常 %d 条 (%.1f%%)",
                           c, n, n / len(out) * 100)
        out[f"__{c}_invalid"] = flag
        flagged_cols += 1
        total_flagged += n
    logger.info("异常值打标完成: %d 列 / 共 %d 处异常",
                flagged_cols, total_flagged)
    return out


# ---------- docx 耐久解析 ----------

# 耐久 docx 真正的数据列名(行索引 4,合并单元格去重后)
DURABILITY_COLUMNS = [
    "目标功率(kW)", "湿度", "温度", "净输出功率(kW)", "电堆电流(A)",
    "平均单体电压(V)", "离均差", "空压机功耗(kW)", "水泵功耗(kW)",
    "冷却水入口温度(℃)", "冷却水出口温度(℃)", "HFR", "LFR", "电压方差",
]


def load_durability_docx(file_paths: Iterable[str]) -> pd.DataFrame:
    """读取多个耐久 docx,按列头解析、拼接为标准 DataFrame。

    每个 docx 视为一段耐久区间(0-5、5-10…)。docx 内 1 个表 65 行:
      - 行 0-3: 元数据(开始/结束时间、系统名称、电堆型号、测试方案、工步数量、测试结果)
      - 行 4:   列头(含合并单元格重复)
      - 行 5-64: 数据行(60 条工步)
    多个 docx 拼接为长表,并附 stage / stage_order / file 列。
    """
    from docx import Document  # 延迟导入,避免无 docx 时影响 CSV 路径

    paths = list(file_paths)
    logger.info("=== docx 耐久解析开始,共 %d 个文件 ===", len(paths))
    t_start = time.perf_counter()

    parts: list[pd.DataFrame] = []
    for i, fp in enumerate(paths):
        p = Path(fp)
        logger.info("[%d/%d] 处理 docx: %s", i + 1, len(paths), p.name)
        try:
            doc = Document(str(p))
            logger.debug("  → Document 加载成功,段落 %d / 表格 %d",
                         len(doc.paragraphs), len(doc.tables))
        except Exception as e:
            logger.error("  → docx 加载失败: %s", e, exc_info=True)
            continue
        # 解析阶段:用正则从文件名抽耐久区间(如 "0-5" / "40-45")
        m = re.search(r"耐久(\d+)-(\d+)", p.stem)
        if m:
            stage_label = f"{m.group(1)}-{m.group(2)}"
            stage_start = int(m.group(1))
            logger.debug("  → 阶段: %s (起始 %d h)", stage_label, stage_start)
        else:
            stage_label = p.stem
            stage_start = 0
            logger.warning("  → 文件名未匹配耐久区间正则,使用 stem=%s", p.stem)

        if not doc.tables:
            logger.warning("  → docx 无表格,跳过")
            continue
        table = doc.tables[0]
        logger.debug("  → 表 0: %d 行 × %d 列",
                     len(table.rows), len(table.columns))

        # 列头行(行 4):合并单元格去重
        if len(table.rows) < 5:
            logger.error("  → 表行数不足 5,无法取列头(实际 %d 行)", len(table.rows))
            continue
        header_cells = [c.text.strip() for c in table.rows[4].cells]
        logger.debug("  → 原始列头(去重前): %s", header_cells)
        # 保序去重
        seen: set[str] = set()
        headers: list[str] = []
        for c in header_cells:
            if c and c not in seen:
                seen.add(c)
                headers.append(c)
        logger.debug("  → 去重后列头: %d 个 %s", len(headers), headers)
        # 若解析出的列数与预期不一致,用标准列名兜底
        if len(headers) != len(DURABILITY_COLUMNS):
            logger.warning("  → 列数与预期不符(实际 %d / 预期 %d),使用标准列名兜底",
                           len(headers), len(DURABILITY_COLUMNS))
            # 按实际列数截取标准列名,避免 DataFrame 列数不匹配 ValueError
            headers = DURABILITY_COLUMNS[:len(headers)]
        else:
            # 解析出的列数正确,但用 docx 全角括号 → 统一映射到标准列名
            headers = DURABILITY_COLUMNS

        # 数据行 5..末
        data_rows: list[list[str]] = []
        for r in table.rows[5:]:
            data_rows.append([c.text.strip() for c in r.cells[:len(headers)]])
        logger.debug("  → 读取 %d 条数据行", len(data_rows))

        # 防御:每行截到 headers 长度,避免行内格子数与列头不一致导致 ValueError
        data_rows = [row[:len(headers)] for row in data_rows]
        df = pd.DataFrame(data_rows, columns=headers)
        # 数值化(统计转换失败的格子)
        for c in df.columns:
            before = df[c].isna().sum()
            df[c] = pd.to_numeric(df[c], errors="coerce")
            new_na = df[c].isna().sum() - before
            if new_na:
                logger.warning("  → 列 %s 数值化失败 %d 格", c, new_na)
        df["stage"] = stage_label
        df["stage_start_h"] = stage_start
        df["file"] = p.name
        # 行内序号(同一阶段第几条工步)
        df["step_idx"] = range(len(df))
        parts.append(df)
        logger.info("  → 入库 %d 行,阶段=%s", len(df), stage_label)

    if not parts:
        logger.warning("=== docx 解析结束,无可用数据 ===")
        return pd.DataFrame()
    out = pd.concat(parts, ignore_index=True)
    out = out.sort_values(["stage_start_h", "step_idx"]).reset_index(drop=True)
    elapsed = time.perf_counter() - t_start
    logger.info(
        "=== docx 解析结束: %d 份 / %d 行 / %d 阶段 / 耗时 %.2fs ===",
        len(parts), len(out), out["stage"].nunique(), elapsed,
    )
    return out


def load_durability_metadata(file_paths: Iterable[str]) -> pd.DataFrame:
    """读取 docx 前 4 行元数据(开始/结束时间、系统名称等)。

    返回 DataFrame,每行一份 docx。
    """
    from docx import Document
    paths = list(file_paths)
    logger.info("读取 docx 元数据: %d 个文件", len(paths))
    rows: list[dict] = []
    for i, fp in enumerate(paths):
        p = Path(fp)
        try:
            doc = Document(str(p))
        except Exception as e:
            logger.error("[%d/%d] %s 加载失败: %s", i + 1, len(paths), p.name, e)
            continue
        if not doc.tables:
            logger.warning("%s 无表格", p.name)
            continue
        t = doc.tables[0]
        m = re.search(r"耐久(\d+)-(\d+)", p.stem)
        stage_label = f"{m.group(1)}-{m.group(2)}" if m else p.stem
        row: dict = {"file": p.name, "stage": stage_label}
        # 元数据键值对:r0-r3 每行第 1 列(键)、第 4 列(值)
        for ri in range(min(4, len(t.rows))):
            cells = [c.text.strip() for c in t.rows[ri].cells]
            if len(cells) >= 4 and cells[0] and cells[3]:
                row[cells[0]] = cells[3]
        rows.append(row)
    return pd.DataFrame(rows)


def peek_docx_structure(file_path: str, max_tables: int = 3, max_rows: int = 5) -> str:
    """打印 docx 结构(段落数/表格数/前 N 行内容),用于确认列名。"""
    from docx import Document

    doc = Document(file_path)
    out_lines = [
        f"== {Path(file_path).name} ==",
        f"段落数: {len(doc.paragraphs)}",
        f"表格数: {len(doc.tables)}",
    ]
    for i, t in enumerate(doc.tables[:max_tables]):
        out_lines.append(f"-- 表 {i}: {len(t.rows)} 行 × {len(t.columns)} 列 --")
        for r in t.rows[:max_rows]:
            cells = [c.text.strip()[:20] for c in r.cells]
            out_lines.append(" | ".join(cells))
    return "\n".join(out_lines)
